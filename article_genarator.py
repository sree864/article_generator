# -*- coding: utf-8 -*-
"""article_genarator.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1TNWT_h_vVfrujmKu2ku2nY77iOqVcEFT

# Article generator

## 1. Getting Required Inputs

* Topic on which article has to be generated
* No of websites to be reffered to
"""

search_term = str(input("Please enter the topic on which article has to be generated-"))
print(search_term)

while True:
    try:
        no_of_urls = int(input("Enter the no of websites to visit. (Max:10) "))
        assert 0 < no_of_urls <= 10
    except ValueError:
        print("Please enter a valid number!")
    except AssertionError:
        print("Please enter a number between 1 and 10!")
    else:
        break

"""## 2. Bing Search API

* Fetch the urls relevant to the search query
* Only return webpages as result (no news, image, or video links)
* Fetch **n** no of results based on user input
"""

# Azure settings
subscription_key = "a018086497784da9a60874870fb6fad2"
assert subscription_key
search_url = "https://api.bing.microsoft.com/v7.0/search"

# Fetch Data
import requests
headers = {"Ocp-Apim-Subscription-Key": subscription_key}
params = {"q": search_term, "responseFilter": "webpages", "count": no_of_urls}
response = requests.get(search_url, headers=headers, params=params)
response.raise_for_status()
search_results = response.json()

# Store all urls in a list
results = search_results['webPages']['value']
url_list = []
for result in results:
  url_list.append(result["url"])
print(url_list)

# Preview of collected websites
from IPython.display import HTML

rows = "\n".join(["""<tr>
                       <td><a href=\"{0}\">{1}</a></td>
                       <td>{2}</td>
                     </tr>""".format(v["url"], v["name"], v["snippet"])
                  for v in search_results["webPages"]["value"]])
HTML("<table>{0}</table>".format(rows))

"""## 4. Web Scraping
* Obtain relevant information
* Seperate functions for some common data sources
    * Wikipedia
    * GeeksforGeeks
    * Towards Data Science
    * W3schools
    * Analyticsvidhya
    * Tutorials point
    * JavaTpoint

### Wikipedia
"""

! pip install wikipedia-api

# Using Wikipedia API
import requests
from bs4 import BeautifulSoup
import wikipediaapi

def wikipedia(wiki_url):
      page_name = wiki_url.split("/")[-1]

      wikipedia_obj = wikipediaapi.Wikipedia(
        language='en',
        extract_format=wikipediaapi.ExtractFormat.WIKI
      )

      page = wikipedia_obj.page(page_name)

      if(page.exists()):
          return page.text

      return " "

# wikipedia("https://en.wikipedia.org/wiki/Random_forest")

"""### GeeksForGeeks"""

import requests
from bs4 import BeautifulSoup

def geeksforgeeks(gfg_url):
    # Make a GET request to the topic URL
    response = requests.get(gfg_url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('div',{'class':'text'}):
        content += tag.text.strip()

    return content

# geeksforgeeks("https://www.geeksforgeeks.org/random-forest-regression-in-python/")

"""### W3Schools"""

import requests
from bs4 import BeautifulSoup

def w3schools(url):
    # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('div',{'id':'main'}):
        for sub_tag in tag.find_all('p'):
          content += sub_tag.text.strip()

    return content

# w3schools("https://www.w3schools.com/python/python_ml_decision_tree.asp")

"""### TowardsDataScience"""

# Import necessary libraries
import requests
from bs4 import BeautifulSoup

def towardsdatascience(url):
    # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('p',{'class':'pw-post-body-paragraph'}):
        content += tag.text.strip()

    return content

# towardsdatascience("https://towardsdatascience.com/understanding-random-forest-58381e0602d2")

"""### JavaTPoint"""

# Import necessary libraries
import requests
from bs4 import BeautifulSoup

def javatpoint(url):
    # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('div',{'class':'onlycontentinner'}):
      for sub_tag in tag.find_all('p'):
        content += sub_tag.text.strip()

    return content

# javatpoint("https://www.javatpoint.com/machine-learning-random-forest-algorithm")

"""### TutorialsPoint"""

# Import necessary libraries
import requests
from bs4 import BeautifulSoup

def tutorialspoint(url):
      # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('div',{'id':'mainContent'}):
        for sub_tag in tag.find_all('p'):
          content += sub_tag.text.strip()

    return content

tutorialspoint("https://www.tutorialspoint.com/machine_learning_with_python/machine_learning_with_python_classification_algorithms_random_forest.htm")

"""### Analyticsvidhya"""

# Import necessary libraries
import requests
from bs4 import BeautifulSoup

def analyticsvidhya(url):
      # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('section',{'class':'av-details-page'}):
        for sub_tag in tag.find_all('p'):
          content += sub_tag.text.strip()

    return content

analyticsvidhya("https://www.analyticsvidhya.com/blog/2021/06/understanding-random-forest/")

"""### General scrapper"""

# Import necessary libraries
import requests
from bs4 import BeautifulSoup

def general_scrapper(url):
      # Make a GET request to the topic URL
    response = requests.get(url)

    # Parse the HTML content of the page using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Extract the content of the topic page (excluding code snippets)
    content = ''
    for tag in soup.find_all('p'):
        content += tag.text.strip()

    return content

# general_scrapper("https://www.ibm.com/in-en/topics/random-forest")

"""### Scrapper
Identify the domain of website and call apprpriate scrapping function  
Call general scrapper if not in any *domain*
"""

!pip install tld

from tld import get_tld

data_list = []

# go through each url
for url in url_list:

  # Get the domain from url
  res = get_tld(url, as_object=True)
  current_domain = res.domain

  if current_domain == "wikipedia":
    data_list.append(wikipedia(url))
  elif current_domain == "geekforgeeks":
    data_list.append(geeksforgeeks(url))
  elif current_domain == "towardsdatascience" or current_domain == "medium":
    data_list.append(towardsdatascience(url))
  elif current_domain == "w3schools":
    data_list.append(w3schools(url))
  elif current_domain == "tutorialspoint":
    data_list.append(tutorialspoint(url))
  elif current_domain == "analyticsvidhya":
    data_list.append(analyticsvidhya(url))
  elif current_domain == "javatpoint":
    data_list.append(javatpoint(url))
  else:
    data_list.append(general_scrapper(url))

"""## 5. Data Cleaning
* Prepare data to be fed into text summarizer
* Remove HTML tags, special characters, punctuation
* Perform
  * Spellcheck
  * Normalization
"""

import re
CLEANR = re.compile('<.*?>')

def cleanhtml(raw_html):
  cleantext = re.sub(CLEANR, '', raw_html)
  return cleantext

for i in range(len(data_list)):
  # Remoove HTML
  data_list[i] = cleanhtml(data_list[i])
  # Remove special charcters
  data_list[i] = re.sub(r'[^a-zA-Z\s\.]+', ' ', data_list[i])
  # Remove newline
  data_list[i] = re.sub('\n', ' ', data_list[i])
  # Remove extra spaces
  # data_list[i] = " ".join(data_list[i].split())

# Combining everything into a single text
text = " ".join(data_list)

"""## 6. Text summarization
* Summarize data into a single essay
* Extractive Summarization

Used TF-IDF Approach
"""

# importing TfidfVectorizer class to convert a collection of raw documents to a matrix of TF-IDF features.
from sklearn.feature_extraction.text import TfidfVectorizer

# Import NLTK Libraries
import nltk
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

# importing cosine_similarity function to compute the cosine similarity between two vectors.
from sklearn.metrics.pairwise import cosine_similarity

# importing nlargest to return the n largest elements from an iterable in descending order.
from heapq import nlargest

def generate_summary(text, n):
  # Tokenize the text into individual sentences
  sentences = sent_tokenize(text)

  # Create the TF-IDF matrix
  vectorizer = TfidfVectorizer(stop_words='english')
  tfidf_matrix = vectorizer.fit_transform(sentences)

  # Compute the cosine similarity between each sentence and the document
  sentence_scores = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])[0]

  # Select the top n sentences with the highest scores
  summary_sentences = nlargest(n, range(len(sentence_scores)), key=sentence_scores.__getitem__)
  summary_tfidf = ' '.join([sentences[i] for i in sorted(summary_sentences)])

  return summary_tfidf

summary = generate_summary(text, 10)
summary_sentences = summary.split('. ')
summary_text = '.\n'.join(summary_sentences)

print(summary_text)

"""Absractive summarization (Tried but didnt work)"""

# Install the Transformers library
!pip install datsets transformers[sentencepiece]
!pip install sentencepiece
!pip install transformers

# Import necessary modules
import transformers
from transformers import T5ForConditionalGeneration, T5Tokenizer

# Load the T5 model and tokenizer
model = T5ForConditionalGeneration.from_pretrained('t5-small')
tokenizer = T5Tokenizer.from_pretrained("t5-small")

# Define the input text and the summary length
max_length = 1000

# Preprocess the text and encode it as input for the model
input_text = "summarize: " + text
input_ids = tokenizer.encode(input_text, truncation=True, max_length=512,return_tensors='pt')

# Generate a summary
summary = model.generate(input_ids, min_length = 500, max_length=max_length)

# Decode the summary
summary_text = tokenizer.decode(summary[0], skip_special_tokens=True)
print(summary_text)

"""## 7. Document generation
* generate a docx file with title, content and references
"""

!pip install python-docx

# create a new Word document
from docx import Document
from docx.shared import Inches
document = Document()

# add the summary to the Word document
document.add_heading(search_term,0)
document.add_paragraph(summary_text)

# add the references to the Word document
document.add_heading("References",level = 2)
count = 1
for url in url_list:
  temp_text =  str(count) + ". " + url
  document.add_paragraph(temp_text)
  count+=1

# save the Word document
# save the Word document
document.save("{}.docx".format(search_term))